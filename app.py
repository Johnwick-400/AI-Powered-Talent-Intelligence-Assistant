import streamlit as st
import requests
import json
import re
import pandas as pd
from datetime import datetime
import os
import tempfile
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib import colors

#dependencies
MISTRAL_API_KEY = ""
GEMINI_API_KEY = ""
MISTRAL_API_URL = "https://api.mistral.ai/v1/chat/completions"
GEMINI_API_URL = (
    "https://generativelanguage.googleapis.com/v1beta/models/"
    "gemini-2.5-pro:generateContent"
)
#ui
st.set_page_config(
    page_title="Your Personal AI Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)
                      
def initialise_state():
    if "jd_data" not in st.session_state:
        st.session_state.jd_data = {
            "job_requirements": {
                "job_title": "",
                "employment": "",
                "role": "",
                "years_experience": "",
                "num_positions": "",
                "ctc_in_inr": "",
                "education": "",
                "work_location": "",
                "mode_of_work": "",
                "interview_rounds": "",
                "salaries_paid_on": "",
                "role_reporting_to": "",
                "must_have": [],
                "good_to_have": [],
            },
            "employer_details": {
                "employer_name": "",
                "Company_name": "",
                "Company_website": "",
            },
            "analysis": {
                "market_analysis": {},
                "role_mapping": [],
            },
        }
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "current_step" not in st.session_state:
        st.session_state.current_step = "start"
    if "missing_fields" not in st.session_state:
        st.session_state.missing_fields = []
    if "current_field_index" not in st.session_state:
        st.session_state.current_field_index = 0
    if "basic_info_index" not in st.session_state:
        st.session_state.basic_info_index = 0
    if "skills_generated" not in st.session_state:
        st.session_state.skills_generated = False
    if "show_analysis" not in st.session_state:
        st.session_state.show_analysis = False
    if "show_mapping" not in st.session_state:
        st.session_state.show_mapping = False
    if "awaiting_user_input" not in st.session_state:
        st.session_state.awaiting_user_input = False

initialise_state()

def extract_text_from_file(uploaded_file):
    try:
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            return "".join(page.extract_text() for page in reader.pages)
        if (
            uploaded_file.type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            doc = docx.Document(uploaded_file)
            return "\n".join(p.text for p in doc.paragraphs)
        try:
            return uploaded_file.read().decode("utf-8")
        except UnicodeDecodeError:
            return uploaded_file.read().decode("latin-1")
    except Exception as exc:
        st.error(f"Error extracting text: {exc}")
        return ""

def call_mistral(prompt: str, max_tokens: int = 1500):
    try:
        response = requests.post(
            MISTRAL_API_URL,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {MISTRAL_API_KEY}",
            },
            json={
                "model": "mistral-small-latest",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": max_tokens,
            },
            timeout=30,
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except Exception as exc:
        st.error(f"Mistral API error: {exc}")
        return None

def call_gemini(prompt: str):
    try:
        response = requests.post(
            f"{GEMINI_API_URL}?key={GEMINI_API_KEY}",
            headers={"Content-Type": "application/json"},
            json={
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {
                    "temperature": 0.7,
                    "maxOutputTokens": 3500
                }
            },
            timeout=60,
        )
        response.raise_for_status()
        data = response.json()
        # Defensive parsing
        try:
            return data["candidates"][0]["content"]["parts"][0]["text"]
        except (KeyError, IndexError, TypeError):
            st.error(f"Gemini API error: Unexpected response format: {data}")
            return None
    except Exception as exc:
        st.error(f"Gemini API error: {exc}")
        return None
def analyse_uploaded_jd(jd_text: str, job_title: str):
    prompt = f"""
You are an AI Assistant, an expert HR analyst. Extract structured information from the job description below and identify which attributes are missing. 
Respond with strict JSON that matches exactly the following schema:

{{
  "job_requirements": {{
    "job_title": "{job_title}",
    "employment": "",
    "role": "",
    "years_experience": "",
    "num_positions": "",
    "ctc_in_inr": "",
    "education": "",
    "work_location": "",
    "mode_of_work": "",
    "interview_rounds": "",
    "salaries_paid_on": "",
    "role_reporting_to": "",
    "must_have": [],
    "good_to_have": []
  }},
  "employer_details": {{
    "employer_name": "",
    "Company_name": "",
    "Company_website": ""
  }},
  "missing_fields": [],
  "extracted_successfully": []
}}

Instructions:
- Extract all available information from the job description
- For missing fields, add them to the "missing_fields" array
- For skills, extract both technical and soft skills
- Ensure all fields are properly populated from the text

Job description:
\"\"\"{jd_text}\"\"\"
"""
    raw = call_mistral(prompt)
    if not raw:
        return {"error": "Failed to analyze job description"}
    
    # Clean JSON response
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.lstrip().startswith("json"):
            raw = "\n".join(raw.splitlines()[1:])
    
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"error": "Failed to parse job description analysis"}

def suggest_skills(job_title, role_value, experience_value, existing_must_have=None, existing_good_to_have=None):
    must_have_existing = existing_must_have or []
    good_to_have_existing = existing_good_to_have or []
    
    prompt = f"""
You are  AI Assistant, a senior technical recruiter. Generate comprehensive skill recommendations for the following position.

Job Title: {job_title}
Role: {role_value}
Experience Required: {experience_value}

Current Must-Have Skills: {must_have_existing}
Current Good-to-Have Skills: {good_to_have_existing}

Generate 8-10 must-have skills and 4-6 good-to-have skills. Avoid duplicates with existing skills.
Focus on:
- Technical skills relevant to the role
- Industry-standard tools and technologies
- Soft skills essential for the position
- Certifications or qualifications

Return JSON format:
{{
  "must_have": [
    "skill1",
    "skill2",
    ...
  ],
  "good_to_have": [
    "skill1",
    "skill2",
    ...
  ],
  "explanation": "Brief explanation of why these skills are recommended"
}}
"""
    
    raw = call_mistral(prompt, max_tokens=1000)
    if not raw:
        
        fallback_skills = get_fallback_skills(job_title, role_value)
        return fallback_skills
    
    
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.lstrip().startswith("json"):
            raw = "\n".join(raw.splitlines()[1:])
    
    try:
        result = json.loads(raw)

        if not result.get("must_have"):
            fallback = get_fallback_skills(job_title, role_value)
            result["must_have"] = fallback["must_have"]
        if not result.get("good_to_have"):
            fallback = get_fallback_skills(job_title, role_value)
            result["good_to_have"] = fallback["good_to_have"]
        return result
    except json.JSONDecodeError:
        return get_fallback_skills(job_title, role_value)

def get_fallback_skills(job_title, role_value):
    """Fallback skills based on common job titles"""
    job_lower = job_title.lower()
    role_lower = role_value.lower()
    
    
    if any(term in job_lower for term in ["developer", "engineer", "programmer"]):
        if "python" in job_lower or "backend" in role_lower:
            return {
                "must_have": ["Python", "Django/Flask", "SQL", "REST APIs", "Git", "Problem Solving", "Code Review", "Testing"],
                "good_to_have": ["Docker", "AWS", "Redis", "Microservices"],
                "explanation": "Essential backend development skills for Python developers"
            }
        elif "javascript" in job_lower or "frontend" in role_lower:
            return {
                "must_have": ["JavaScript", "React/Vue/Angular", "HTML/CSS", "Git", "Responsive Design", "Problem Solving", "Code Review", "Testing"],
                "good_to_have": ["TypeScript", "Node.js", "Webpack", "GraphQL"],
                "explanation": "Core frontend development skills"
            }
        elif "java" in job_lower:
            return {
                "must_have": ["Java", "Spring Framework", "SQL", "REST APIs", "Git", "Problem Solving", "Code Review", "Testing"],
                "good_to_have": ["Spring Boot", "Docker", "Microservices", "Kafka"],
                "explanation": "Essential Java development skills"
            }
    
    elif any(term in job_lower for term in ["data", "analyst", "scientist"]):
        return {
            "must_have": ["Python/R", "SQL", "Data Analysis", "Statistics", "Excel", "Data Visualization", "Problem Solving", "Communication"],
            "good_to_have": ["Machine Learning", "Tableau/Power BI", "Big Data", "Apache Spark"],
            "explanation": "Core data analysis and science skills"
        }
    
    elif any(term in job_lower for term in ["manager", "lead", "senior"]):
        return {
            "must_have": ["Leadership", "Project Management", "Communication", "Team Management", "Strategic Planning", "Problem Solving", "Decision Making", "Mentoring"],
            "good_to_have": ["Agile/Scrum", "Budget Management", "Stakeholder Management", "Process Improvement"],
            "explanation": "Essential leadership and management skills"
        }
    
    
    return {
        "must_have": ["Communication", "Problem Solving", "Teamwork", "Time Management", "Adaptability", "Critical Thinking", "Attention to Detail", "Professional Skills"],
        "good_to_have": ["Leadership", "Project Management", "Technical Writing", "Analytical Skills"],
        "explanation": "General professional skills applicable to most roles"
    }

def perform_role_mapping(jd_data):
    req = jd_data["job_requirements"]
    prompt = f"""
You are  AI Assistant, an experienced recruitment analyst. Map the following job description to alternative role titles that candidates might use in their profiles or resumes.

Job Title: {req.get('job_title', '')}
Role Description: {req.get('role', '')}
Experience Level: {req.get('years_experience', '')}
Must-Have Skills: {', '.join(req.get('must_have', []))}
Good-to-Have Skills: {', '.join(req.get('good_to_have', []))}

Provide 6-8 alternative role titles with match percentage and reasoning.

Return JSON format:
{{
  "potential_roles": [
    {{
      "role_title": "Alternative Role Title",
      "match_percentage": 85,
      "reasoning": "Why this role matches"
    }}
  ]
}}
"""

    raw = call_mistral(prompt, max_tokens=1000)
    if not raw:
        return {"potential_roles": []}

    # Extract JSON from markdown/code block if present
    json_match = re.search(r'\{[\s\S]*\}', raw)
    if json_match:
        raw = json_match.group()

    try:
        return json.loads(raw)
    except Exception:
        st.error(f"Role mapping: Could not parse response: {raw}")
        return {"potential_roles": []}
def analyze_job_market(jd_data):
    req = jd_data["job_requirements"]
    ctc_range = req.get('ctc_in_inr', '')
    min_ctc = max_ctc = ""
    if ctc_range:
        ctc_match = re.search(r'(\d+)\s*-\s*(\d+)', ctc_range)
        if ctc_match:
            min_ctc = ctc_match.group(1)
            max_ctc = ctc_match.group(2)
        else:
            min_ctc = max_ctc = ctc_range

    jd_text = f"""
Job Title: {req.get('job_title', '')}
Role Description: {req.get('role', '')}
Experience Required: {req.get('years_experience', '')}
Location: {req.get('work_location', '')}
Employment Type: {req.get('employment', '')}
Must-Have Skills: {', '.join(req.get('must_have', []))}
Good-to-Have Skills: {', '.join(req.get('good_to_have', []))}
Education: {req.get('education', '')}
"""

    prompt = f"""
You are an expert IT Recruiter in India. Analyze the following job description and provide market analysis.

Job Description:
{jd_text}

Analysis Parameters:
- Salary Range: {min_ctc}-{max_ctc} LPA

Tasks:
1. Extract the mandatory must-have technical skills from the JD
2. Analyze the Indian job market for immediate joiners with these skills
3. Provide realistic market data based on current hiring trends

Create skill combinations (starting with the most important skill, then adding others progressively) and analyze the market for each combination.

Reply ONLY in this JSON format (no explanation, no markdown, no extra text):
{{
  "extracted_skills": ["skill1", "skill2", "skill3"],
  "market_analysis": [
    {{
      "skills": "skill1",
      "num_candidates": 1500,
      "experience_range": "2-5 years",
      "salary_range": "8-15 LPA"
    }},
    {{
      "skills": "skill1 + skill2",
      "num_candidates": 850,
      "experience_range": "2-5 years",
      "salary_range": "10-18 LPA"
    }},
    {{
      "skills": "skill1 + skill2 + skill3",
      "num_candidates": 300,
      "experience_range": "3-6 years",
      "salary_range": "12-20 LPA"
    }}
  ],
  "overall_analysis": {{
    "market_demand": "High/Medium/Low",
    "candidate_availability": "High/Medium/Low",
    "salary_competitiveness": "Above Market/Market Rate/Below Market",
    "time_to_hire": "2-3 weeks",
    "competition_level": "High/Medium/Low"
  }},
  "recommendations": [
    "Recommendation 1",
    "Recommendation 2",
    "Recommendation 3"
  ]
}}

IMPORTANT: Output ONLY valid JSON. Do not include any explanation, markdown, or extra text.
"""

    raw = call_gemini(prompt)
    # st.write("Gemini raw output:", raw)  # For debugging
    if not raw:
        return get_fallback_market_analysis_detailed(req)
    parsed = clean_and_parse_json(raw)
    if parsed and parsed.get("market_analysis"):
        return parsed

    # If all fails, fallback
    return get_fallback_market_analysis_detailed(req)


def clean_and_parse_json(text_response: str):
    """
    Robustly cleans and parses a JSON string from an LLM response.
    Handles markdown code blocks, trailing commas, and other common issues.
    """
    json_match = re.search(r'```json\s*(\{[\s\S]*\})\s*```|(\{[\s\S]*\})', text_response, re.DOTALL)
    if not json_match:
        st.error("❌ No valid JSON structure found in the API response.")
        st.write("Raw Response:", text_response)
        return None
    json_str = json_match.group(1) or json_match.group(2)
    try:
        json_str = re.sub(r',\s*([}\]])', r'\1', json_str)
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        st.error(f"❌ Failed to decode JSON. Error: {e}")
        st.write("Problematic JSON String:", json_str)
        return None
def create_fallback_skill_analysis(req):
    """Create fallback skill-based analysis"""
    must_have = req.get('must_have', [])
    if not must_have:
        return []
    
    analysis = []
    base_candidates = 2000
    
    for i, skill in enumerate(must_have[:5]):  
        if i == 0:
            
            analysis.append({
                "skills": skill,
                "num_candidates": base_candidates,
                "experience_range": req.get('years_experience', '2-5 years'),
                "salary_range": req.get('ctc_in_inr', '8-15 LPA')
            })
        else:
            skills_combo = " + ".join(must_have[:i+1])
            candidates = max(100, base_candidates // ((i+1) * 2))
            
            analysis.append({
                "skills": skills_combo,
                "num_candidates": candidates,
                "experience_range": req.get('years_experience', '2-5 years'),
                "salary_range": req.get('ctc_in_inr', '8-15 LPA')
            })
    
    return analysis

def get_fallback_market_analysis_detailed(req):
    """Enhanced fallback market analysis with skill combinations"""
    must_have = req.get('must_have', [])
    
    return {
        "extracted_skills": must_have[:5] if must_have else ["General Skills"],
        "market_analysis": create_fallback_skill_analysis(req),
        "overall_analysis": {
            "market_demand": "Medium",
            "candidate_availability": "Medium",
            "salary_competitiveness": "Market Rate",
            "time_to_hire": "3-4 weeks",
            "competition_level": "Medium"
        },
        "recommendations": [
            "Consider highlighting company culture and growth opportunities",
            "Streamline the interview process to reduce time-to-hire",
            "Offer competitive benefits package along with salary",
            "Focus on candidates with core skills and train for additional requirements"
        ]
    }

def enhance_user_input(user_input, field_name):
    """Enhance user input using AI to create more professional JD content"""
    if not user_input or len(user_input.strip()) < 3:
        return user_input
    
    prompt = f"""
You are  AI Assistant. Enhance the following user input for the "{field_name}" field in a job description. 
Make it more professional and detailed while preserving the original intent.

User input: "{user_input}"
Field: {field_name}

Return only the enhanced text, nothing else.
"""
    
    enhanced = call_mistral(prompt, max_tokens=200)
    return enhanced if enhanced else user_input

def write_docx(jd_data, path):
    try:
        document = docx.Document()
        
       
        header = document.add_heading("Job Description", 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        company = jd_data["employer_details"].get("Company_name", "")
        if company:
            company_para = document.add_paragraph(company)
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            company_para.runs[0].bold = True
            company_para.runs[0].font.size = Pt(14)
        
        
        job_title = jd_data["job_requirements"].get("job_title", "Position")
        title_heading = document.add_heading(job_title, 1)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_heading("Job Requirements", 2)
        
        field_labels = {
            "employment": "Employment Type",
            "role": "Role Description",
            "years_experience": "Experience Required",
            "num_positions": "Number of Positions",
            "ctc_in_inr": "CTC Range (INR)",
            "education": "Educational Requirements",
            "work_location": "Work Location",
            "mode_of_work": "Work Mode",
            "interview_rounds": "Interview Process",
            "salaries_paid_on": "Salary Payment",
            "role_reporting_to": "Reporting Manager"
        }
        
        for field, label in field_labels.items():
            value = jd_data["job_requirements"].get(field, "")
            if value:
                document.add_paragraph(f"{label}: {value}")
        
        if jd_data["job_requirements"]["must_have"]:
            document.add_heading("Must-Have Skills", 2)
            for skill in jd_data["job_requirements"]["must_have"]:
                document.add_paragraph(skill, style="List Bullet")
        
        if jd_data["job_requirements"]["good_to_have"]:
            document.add_heading("Good-to-Have Skills", 2)
            for skill in jd_data["job_requirements"]["good_to_have"]:
                document.add_paragraph(skill, style="List Bullet")
        
        if any(jd_data["employer_details"].values()):
            document.add_heading("Company Information", 2)
            employer_name = jd_data["employer_details"].get("employer_name", "")
            website = jd_data["employer_details"].get("Company_website", "")
            
            if employer_name:
                document.add_paragraph(f"Employer: {employer_name}")
            if website:
                document.add_paragraph(f"Website: {website}")
      
        document.add_paragraph()
        footer = document.add_paragraph("Generated by  AI Assistant")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].italic = True
        footer.runs[0].font.size = Pt(10)
        
        document.save(path)
        return True
    except Exception as exc:
        st.error(f"Error generating DOCX: {exc}")
        return False

def write_pdf(jd_data, path):
    try:
        # Create document with letter size pages
        doc = SimpleDocTemplate(path, pagesize=letter)
        
        # Get base styles and create custom styles
        styles = getSampleStyleSheet()
        
        # Instead of adding styles to the stylesheet, create new ParagraphStyle instances
        title_style = ParagraphStyle(
            'CustomTitle',
            fontSize=20,
            alignment=TA_CENTER,
            spaceAfter=12,
            parent=styles['Normal']
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=6,
            parent=styles['Normal']
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            fontSize=14,
            alignment=TA_LEFT,
            spaceAfter=6,
            parent=styles['Normal']
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            fontSize=12,
            alignment=TA_JUSTIFY,
            spaceAfter=3,
            parent=styles['Normal']
        )
        
        list_style = ParagraphStyle(
            'CustomList',
            fontSize=12,
            alignment=TA_LEFT,
            leftIndent=20,
            spaceAfter=3,
            parent=styles['Normal']
        )
        
        footer_style = ParagraphStyle(
            'CustomFooter',
            fontSize=10,
            alignment=TA_CENTER,
            textColor=colors.grey,
            parent=styles['Normal']
        )
        
        # Story will hold all elements
        story = []
        
        # Add title
        story.append(Paragraph("Job Description", title_style))
        
        # Add company name if available
        company = jd_data["employer_details"].get("Company_name", "")
        if company:
            story.append(Paragraph(company, heading1_style))
            story.append(Spacer(1, 12))
        
        # Job title
        job_title = jd_data["job_requirements"].get("job_title", "Position")
        story.append(Paragraph(job_title, heading1_style))
        story.append(Spacer(1, 20))
        
        # Job Requirements
        story.append(Paragraph("Job Requirements", heading2_style))
        
        field_labels = {
            "employment": "Employment Type",
            "role": "Role Description",
            "years_experience": "Experience Required",
            "num_positions": "Number of Positions",
            "ctc_in_inr": "CTC Range (INR)",
            "education": "Educational Requirements",
            "work_location": "Work Location",
            "mode_of_work": "Work Mode",
            "interview_rounds": "Interview Process",
            "salaries_paid_on": "Salary Payment",
            "role_reporting_to": "Reporting Manager"
        }
        
        for field, label in field_labels.items():
            value = jd_data["job_requirements"].get(field, "")
            if value:
                story.append(Paragraph(f"<b>{label}:</b> {value}", normal_style))
                story.append(Spacer(1, 5))
        
        # Must-have skills
        if jd_data["job_requirements"]["must_have"]:
            story.append(Paragraph("Must-Have Skills", heading2_style))
            must_have_list = []
            for skill in jd_data["job_requirements"]["must_have"]:
                must_have_list.append(ListItem(Paragraph(skill, list_style)))
            story.append(ListFlowable(must_have_list, bulletType="bullet"))
            story.append(Spacer(1, 10))
        
        # Good-to-have skills
        if jd_data["job_requirements"]["good_to_have"]:
            story.append(Paragraph("Good-to-Have Skills", heading2_style))
            good_to_have_list = []
            for skill in jd_data["job_requirements"]["good_to_have"]:
                good_to_have_list.append(ListItem(Paragraph(skill, list_style)))
            story.append(ListFlowable(good_to_have_list, bulletType="bullet"))
            story.append(Spacer(1, 10))
        
        # Company information
        if any(jd_data["employer_details"].values()):
            story.append(Paragraph("Company Information", heading2_style))
            employer_name = jd_data["employer_details"].get("employer_name", "")
            website = jd_data["employer_details"].get("Company_website", "")
            
            if employer_name:
                story.append(Paragraph(f"<b>Contact:</b> {employer_name}", normal_style))
            if website:
                story.append(Paragraph(f"<b>Website:</b> {website}", normal_style))
            story.append(Spacer(1, 10))
        
        # Footer
        story.append(Spacer(1, 20))
        story.append(Paragraph("Generated by  AI Assistant", footer_style))
        
        # Build document
        doc.build(story)
        return True
    except Exception as exc:
        st.error(f"Error generating PDF: {exc}")
        return False
    
def get_question_prompt(field_name, current_data):
    """Get contextual question prompts for better user experience"""
    job_title = current_data.get("job_title", "")
    
    prompts = {
        "employment": f"What type of employment is this {job_title} position? (e.g., Full-time, Part-time, Contract, Internship)",
        "role": f"Could you describe the main responsibilities and duties for this {job_title} role?",
        "years_experience": f"How many years of experience are required for this {job_title} position? (e.g., 2-5 years, 0-2 years)",
        "num_positions": "How many positions are you looking to fill for this role?",
        "ctc_in_inr": "What's the CTC range for this position in INR? (e.g., 8-12 LPA, 15-20 LPA)",
        "education": "What educational qualifications are required? (e.g., Bachelor's in Computer Science, MBA)",
        "work_location": "What's the work location for this position? (e.g., Bangalore, Mumbai, Remote)",
        "mode_of_work": "What's the work mode? (e.g., Remote, Hybrid, On-site)",
        "interview_rounds": "How many interview rounds will there be? (e.g., 3 rounds, HR + Technical + Final)",
        "salaries_paid_on": "When are salaries paid? (e.g., Monthly, Bi-weekly)",
        "role_reporting_to": "Who will this role report to? (e.g., Engineering Manager, Team Lead, CTO)",
        "employer_name": "What's the hiring manager's name or your name?",
        "Company_name": "What's the company name?",
        "Company_website": "What's the company website? (optional)"
    }
    
    return prompts.get(field_name, f"Could you provide the {field_name.replace('_', ' ')}?")

def display_market_analysis(analysis_data):
    """Display market analysis results"""
    st.subheader("📊 Market Analysis")
    
    if not analysis_data:
        st.warning("No market analysis data available.")
        return
    
    if "market_analysis" in analysis_data and analysis_data["market_analysis"]:
        st.markdown("### 🎯 Skill-Based Candidate Analysis")
        
        df_data = []
        for item in analysis_data["market_analysis"]:
            df_data.append({
                "Skills": item.get("skills", ""),
                "Candidates": item.get("num_candidates", 0),
                "Experience": item.get("experience_range", ""),
                "Salary Range": item.get("salary_range", "")
            })
        
        if df_data:
            df = pd.DataFrame(df_data)
            st.dataframe(df, use_container_width=True)
    
    
    if "overall_analysis" in analysis_data:
        overall = analysis_data["overall_analysis"]
        st.markdown("### 📈 Overall Market Assessment")
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Market Demand", overall.get("market_demand", "N/A"))
            st.metric("Candidate Availability", overall.get("candidate_availability", "N/A"))
            st.metric("Competition Level", overall.get("competition_level", "N/A"))
        
        with col2:
            st.metric("Salary Competitiveness", overall.get("salary_competitiveness", "N/A"))
            st.metric("Expected Time to Hire", overall.get("time_to_hire", "N/A"))
    
   
    if "recommendations" in analysis_data and analysis_data["recommendations"]:
        st.markdown("### 💡 Recommendations")
        for i, rec in enumerate(analysis_data["recommendations"], 1):
            st.markdown(f"{i}. {rec}")


def display_role_mapping(role_mapping_data):
    """Display role mapping results"""
    st.subheader("🔄 Alternative Role Titles")
    
    if not role_mapping_data or "potential_roles" not in role_mapping_data:
        st.warning("No role mapping data available.")
        return
    
    potential_roles = role_mapping_data["potential_roles"]
    if not potential_roles:
        st.info("No alternative roles found.")
        return
    
    for role in potential_roles:
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**{role.get('role_title', 'Unknown Role')}**")
            st.markdown(f"*{role.get('reasoning', 'No reasoning provided')}*")
        with col2:
            match_pct = role.get('match_percentage', 0)
            st.metric("Match", f"{match_pct}%")
        st.divider()

def display_chat_messages():
    """Display chat messages"""
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])

def add_message(role, content):
    """Add a message to chat history"""
    st.session_state.messages.append({"role": role, "content": content})


def start_basic_info_collection():
    """Start collecting basic information"""
    st.session_state.current_step = "collect_basic_info"
    st.session_state.awaiting_user_input = True
    
    
    required_fields = [
        "employment", "role", "years_experience", "num_positions", 
        "ctc_in_inr", "education", "work_location", "mode_of_work",
        "interview_rounds", "salaries_paid_on", "role_reporting_to",
        "employer_name", "Company_name"
    ]
    
   
    missing = []
    for field in required_fields:
        if field in ["employer_name", "Company_name", "Company_website"]:
            # Check in employer_details
            value = st.session_state.jd_data["employer_details"].get(field, "")
            if isinstance(value, str):
                if not value.strip():
                    missing.append(field)
            elif not value:
                missing.append(field)
        else:
            # Check in job_requirements
            try:
                value = st.session_state.jd_data["job_requirements"].get(field, "")
                if isinstance(value, str):
                    if not value.strip():
                        missing.append(field)
                elif isinstance(value, list):
                    if not value:  # empty list
                        missing.append(field)
                elif not value:
                    missing.append(field)
            except KeyError:
                missing.append(field)

    st.session_state.missing_fields = missing
    st.session_state.current_field_index = 0

    if missing:
        field = missing[0]
        question = get_question_prompt(field, st.session_state.jd_data["job_requirements"])
        add_message("assistant", question)
    else:
        st.session_state.current_step = "skill_suggestions"
        add_message("assistant", "Great! All basic information is complete. Let me suggest some relevant skills for this position.")

def handle_basic_info_input(user_input):
    """Handle user input for basic information collection"""
    if not st.session_state.missing_fields:
        return
    
    current_field = st.session_state.missing_fields[st.session_state.current_field_index]
    
    
    enhanced_input = enhance_user_input(user_input, current_field)
    
    if current_field in ["employer_name", "Company_name", "Company_website"]:
        st.session_state.jd_data["employer_details"][current_field] = enhanced_input
    else:
        st.session_state.jd_data["job_requirements"][current_field] = enhanced_input
    
    
    st.session_state.current_field_index += 1
    
    if st.session_state.current_field_index < len(st.session_state.missing_fields):
        
        next_field = st.session_state.missing_fields[st.session_state.current_field_index]
        question = get_question_prompt(next_field, st.session_state.jd_data["job_requirements"])
        add_message("assistant", f"Thank you! {question}")
    else:
        # All fields collected
        st.session_state.current_step = "skill_suggestions"
        st.session_state.awaiting_user_input = False
        add_message("assistant", "Perfect! All basic information is collected. Let me suggest some relevant skills for this position.")

def handle_skill_suggestions():
    """Handle skill suggestions workflow"""
    if not st.session_state.skills_generated:
        req = st.session_state.jd_data["job_requirements"]
        job_title = req.get("job_title", "")
        role = req.get("role", "")
        experience = req.get("years_experience", "")
        
        with st.spinner("Generating skill suggestions..."):
            skills_result = suggest_skills(
                job_title, 
                role, 
                experience,
                req.get("must_have", []),
                req.get("good_to_have", [])
            )
        
        if skills_result:
            # Add new skills to existing ones
            existing_must_have = set(req.get("must_have", []))
            existing_good_to_have = set(req.get("good_to_have", []))
            
            new_must_have = [skill for skill in skills_result.get("must_have", []) 
                           if skill not in existing_must_have]
            new_good_to_have = [skill for skill in skills_result.get("good_to_have", []) 
                              if skill not in existing_good_to_have]
            
            st.session_state.jd_data["job_requirements"]["must_have"].extend(new_must_have)
            st.session_state.jd_data["job_requirements"]["good_to_have"].extend(new_good_to_have)
            
            explanation = skills_result.get("explanation", "")
            add_message("assistant", f"I've suggested additional skills based on the job requirements. {explanation}")
            
        st.session_state.skills_generated = True

# ---------------------------------------------------------------------
# 9. Main Streamlit application
# ---------------------------------------------------------------------
def main():
    st.title("🤖  AI Assistant")
    st.markdown("*Your intelligent companion for creating, analyzing, and optimizing job descriptions*")

    # ---------------- Sidebar: Enhanced summary and progress ----------------
    with st.sidebar:
        st.subheader("📋 Job Description Summary")
        snapshot = st.session_state.jd_data["job_requirements"]
        employer = st.session_state.jd_data["employer_details"]
        
        # Display current job info
        if snapshot.get("job_title"):
            st.markdown(f"**Position:** {snapshot['job_title']}")
        if employer.get("Company_name"):
            st.markdown(f"**Company:** {employer['Company_name']}")
        if snapshot.get("years_experience"):
            st.markdown(f"**Experience:** {snapshot['years_experience']}")
        if snapshot.get("work_location"):
            st.markdown(f"**Location:** {snapshot['work_location']}")
        if snapshot.get("ctc_in_inr"):
            st.markdown(f"**CTC:** {snapshot['ctc_in_inr']}")
        
        # Skills summary
        must_have_count = len(snapshot.get("must_have", []))
        good_to_have_count = len(snapshot.get("good_to_have", []))
        
        if must_have_count > 0:
            st.markdown(f"**Must-have skills:** {must_have_count} skills")
        if good_to_have_count > 0:
            st.markdown(f"**Good-to-have skills:** {good_to_have_count} skills")

        # Progress indicator
        if st.session_state.current_step != "start":
            st.subheader("📊 Progress")
            steps = ["📝 Basic Info", "🎯 Skills", "📈 Analysis", "📄 Results"]
            step_mapping = {
                "collect_basic_info": 0,
                "analyse_uploaded": 0,
                "skill_suggestions": 1,
                "market_analysis": 2,
                "results": 3,
                "download": 3,
            }
            current_step_idx = step_mapping.get(st.session_state.current_step, 0)
            
            for i, step_name in enumerate(steps):
                if i < current_step_idx:
                    st.markdown(f"✅ {step_name}")
                elif i == current_step_idx:
                    st.markdown(f"🔄 {step_name}")
                else:
                    st.markdown(f"⏳ {step_name}")

        # Quick actions
        st.subheader("⚡ Quick Actions")
        if st.session_state.current_step not in ["start"]:
            if st.button("🔄 Start Over", use_container_width=True):
                # Reset all session state
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()

    # Main content area
    col_chat, col_actions = st.columns([3, 1])

    # ---------- Chat column ------------------------------------------------
    with col_chat:
        if st.session_state.current_step == "start":
            st.subheader("🚀 Welcome to  AI Assistant")
            st.markdown("You can either upload an existing job description or create a new one from scratch. Let's get started!")
            tab_upload, tab_create = st.tabs(["📁 Upload Existing JD", "✨ Create New JD"])
            
            with tab_upload:
                st.markdown("Upload your existing job description and I'll analyze and enhance it.")
                
                col1, col2 = st.columns([2, 1])
                with col1:
                    title_input = st.text_input("Job Title", placeholder="e.g., Senior Software Engineer", key="upload_title")
                with col2:
                    st.markdown("")  # Spacing
                
                uploaded_file = st.file_uploader(
                    "Choose a file",
                    type=["pdf", "docx", "txt"],
                    help="Upload PDF, DOCX, or TXT files containing your job description"
                )
                
                if st.button("📊 Analyze Job Description", type="primary", disabled=not (uploaded_file and title_input)):
                    if uploaded_file and title_input:
                        with st.spinner("Analyzing job description..."):
                            jd_text = extract_text_from_file(uploaded_file)
                            if jd_text:
                                analysis_result = analyse_uploaded_jd(jd_text, title_input)
                                
                                if "error" not in analysis_result:
                                    # Update session state with extracted data
                                    st.session_state.jd_data.update(analysis_result)
                                    st.session_state.current_step = "analyse_uploaded"
                                    
                                    add_message("assistant", f"I've analyzed your job description for '{title_input}'. Let me show you what I found and suggest improvements.")
                                    st.rerun()
                                else:
                                    st.error("Failed to analyze the job description. Please try again.")
                            else:
                                st.error("Could not extract text from the uploaded file.")
            
            with tab_create:
                st.markdown("Create a new job description from scratch with AI assistance.")
                
                new_title = st.text_input("Job Title", placeholder="e.g., Python Developer, Product Manager", key="new_title")
                
                if st.button("🚀 Start Creating", type="primary", disabled=not new_title):
                    if new_title:
                        st.session_state.jd_data["job_requirements"]["job_title"] = new_title
                        add_message("assistant", f"Great! Let's create a job description for '{new_title}'. I'll ask you a few questions to gather the necessary information.")
                        start_basic_info_collection()
                        st.rerun()

        else:
            # Display chat interface
            st.subheader("💬 AI Assistant Chat")
            
            # Chat container
            chat_container = st.container(height=400)
            with chat_container:
                display_chat_messages()
            
            # Handle different workflow steps
            if st.session_state.current_step == "analyse_uploaded":
                st.markdown("### 📋 Extracted Information")
                
                # Show extracted data
                req = st.session_state.jd_data["job_requirements"]
                emp = st.session_state.jd_data["employer_details"]
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Job Requirements:**")
                    for field, value in req.items():
                        if value and field != "must_have" and field != "good_to_have":
                            st.markdown(f"- **{field.replace('_', ' ').title()}:** {value}")
                
                with col2:
                    st.markdown("**Company Details:**")
                    for field, value in emp.items():
                        if value:
                            st.markdown(f"- **{field.replace('_', ' ').title()}:** {value}")
                
                # Skills
                if req.get("must_have"):
                    st.markdown("**Must-Have Skills:**")
                    st.write(", ".join(req["must_have"]))
                
                if req.get("good_to_have"):
                    st.markdown("**Good-to-Have Skills:**")
                    st.write(", ".join(req["good_to_have"]))
                
                # Continue button
                if st.button("✅ Continue with Enhancement", type="primary"):
                    start_basic_info_collection()
                    st.rerun()
    
            elif st.session_state.current_step == "collect_basic_info":
                # Handle basic info collection
                if st.session_state.awaiting_user_input and st.session_state.missing_fields:
                    current_field = st.session_state.missing_fields[st.session_state.current_field_index]
                    st.markdown(f"**Current field:** {current_field.replace('_', ' ').title()}")
                    
                    # Progress bar
                    progress = (st.session_state.current_field_index) / len(st.session_state.missing_fields)
                    st.progress(progress, text=f"Progress: {st.session_state.current_field_index}/{len(st.session_state.missing_fields)} fields completed")
                
                # Check if we should auto-proceed to skills
                if not st.session_state.awaiting_user_input and st.session_state.current_step == "collect_basic_info":
                    st.session_state.current_step = "skill_suggestions"
                    st.rerun()
            
            elif st.session_state.current_step == "skill_suggestions":
                handle_skill_suggestions()
                
                # Display current skills
                req = st.session_state.jd_data["job_requirements"]
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Must-Have Skills:**")
                    must_have = req.get("must_have", [])
                    for i, skill in enumerate(must_have):
                        st.markdown(f"{i+1}. {skill}")
                
                with col2:
                    st.markdown("**Good-to-Have Skills:**")
                    good_to_have = req.get("good_to_have", [])
                    for i, skill in enumerate(good_to_have):
                        st.markdown(f"{i+1}. {skill}")
                
                # Options to proceed
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("📊 Run Market Analysis", type="primary"):
                        st.session_state.current_step = "market_analysis"
                        st.rerun()
                
                with col_btn2:
                    if st.button("📄 Generate Final JD", type="secondary"):
                        st.session_state.current_step = "results"
                        st.rerun()
            
            elif st.session_state.current_step == "market_analysis":
                st.markdown("### 📊 Market Analysis & Role Mapping")
                
                # Run analysis if not done
                if not st.session_state.show_analysis:
                    with st.spinner("Analyzing job market..."):
                        analysis_result = analyze_job_market(st.session_state.jd_data)
                        st.session_state.jd_data["analysis"]["market_analysis"] = analysis_result
                        st.session_state.show_analysis = True
                
                if not st.session_state.show_mapping:
                    with st.spinner("Finding alternative role titles..."):
                        mapping_result = perform_role_mapping(st.session_state.jd_data)
                        st.session_state.jd_data["analysis"]["role_mapping"] = mapping_result
                        st.session_state.show_mapping = True
                
                # Display results
                if st.session_state.show_analysis:
                    display_market_analysis(st.session_state.jd_data["analysis"]["market_analysis"])
                
                if st.session_state.show_mapping:
                    display_role_mapping(st.session_state.jd_data["analysis"]["role_mapping"])
                
                # Continue button
                if st.button("📄 Generate Final Job Description", type="primary"):
                    st.session_state.current_step = "results"
                    st.rerun()
            
            elif st.session_state.current_step == "results":
                st.markdown("### 📄 Complete Job Description")
                
                # Display final JD
                req = st.session_state.jd_data["job_requirements"]
                emp = st.session_state.jd_data["employer_details"]
                
                # Job title and company
                st.markdown(f"# {req.get('job_title', 'Job Position')}")
                if emp.get("Company_name"):
                    st.markdown(f"**Company:** {emp['Company_name']}")
                
                # Job details
                st.markdown("## Job Details")
                details = [
                    ("Employment Type", req.get("employment")),
                    ("Experience Required", req.get("years_experience")),
                    ("Number of Positions", req.get("num_positions")),
                    ("CTC Range", req.get("ctc_in_inr")),
                    ("Location", req.get("work_location")),
                    ("Work Mode", req.get("mode_of_work")),
                    ("Education", req.get("education")),
                    ("Reporting To", req.get("role_reporting_to")),
                ]
                
                for label, value in details:
                    if value:
                        st.markdown(f"**{label}:** {value}")
                
                # Role description
                if req.get("role"):
                    st.markdown("## Role Description")
                    st.markdown(req["role"])
                
                # Skills
                if req.get("must_have"):
                    st.markdown("## Must-Have Skills")
                    for skill in req["must_have"]:
                        st.markdown(f"- {skill}")
                
                if req.get("good_to_have"):
                    st.markdown("## Good-to-Have Skills")
                    for skill in req["good_to_have"]:
                        st.markdown(f"- {skill}")
                
                # Interview process
                if req.get("interview_rounds"):
                    st.markdown("## Interview Process")
                    st.markdown(req["interview_rounds"])
                
                # Company info
                if emp.get("employer_name") or emp.get("Company_website"):
                    st.markdown("## Company Information")
                    if emp.get("employer_name"):
                        st.markdown(f"**Contact:** {emp['employer_name']}")
                    if emp.get("Company_website"):
                        st.markdown(f"**Website:** {emp['Company_website']}")
                
                st.session_state.current_step = "download"
            
            # User input for chat
            if st.session_state.awaiting_user_input:
                user_input = st.chat_input("Type your response here...")
                if user_input:
                    add_message("user", user_input)
                    
                    if st.session_state.current_step == "collect_basic_info":
                        handle_basic_info_input(user_input)
                    
                    st.rerun()

    # ---------- Actions column ---------------------------------------------
    with col_actions:
        st.subheader("⚡ Actions")
        
        if st.session_state.current_step == "download" or st.session_state.current_step == "results":
            st.markdown("### 📥 Download Options")
            
            # Generate DOCX
            if st.button("📄 Download DOCX", use_container_width=True):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                    file_path = tmp_file.name
                    if write_docx(st.session_state.jd_data, file_path):
                        tmp_file.close()  # Ensure file is closed before reading
                        with open(file_path, "rb") as f:
                            st.download_button(
                                label="📄 Download Job Description",
                                data=f.read(),
                                file_name=f"{st.session_state.jd_data['job_requirements'].get('job_title', 'job_description').replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                            os.unlink(file_path)
            
            # Generate PDF
            if st.button("📝 Download PDF", use_container_width=True):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    file_path = tmp_file.name
                    if write_pdf(st.session_state.jd_data, file_path):
                        tmp_file.close()  # Ensure file is closed before reading
                        with open(file_path, "rb") as f:
                            st.download_button(
                                label="📝 Download Job Description (PDF)",
                                data=f.read(),
                                file_name=f"{st.session_state.jd_data['job_requirements'].get('job_title', 'job_description').replace(' ', '_')}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                            os.unlink(file_path)
            
            # Generate JSON
            if st.button("📊 Download JSON", use_container_width=True):
                json_data = json.dumps(st.session_state.jd_data, indent=2)
                st.download_button(
                    label="📊 Download Data",
                    data=json_data,
                    file_name=f"{st.session_state.jd_data['job_requirements'].get('job_title', 'job_description').replace(' ', '_')}.json",
                    mime="application/json",
                    use_container_width=True
                )
        
        # Show analysis options
        if st.session_state.current_step in ["skill_suggestions", "results", "download"]:
            st.markdown("### 📊 Analysis")
            
            if not st.session_state.show_analysis:
                if st.button("🔍 Market Analysis", use_container_width=True):
                    st.session_state.current_step = "market_analysis"
                    st.rerun()
            
            if not st.session_state.show_mapping:
                if st.button("🔄 Role Mapping", use_container_width=True):
                    st.session_state.current_step = "market_analysis"
                    st.rerun()
        
        # Navigation
        if st.session_state.current_step not in ["start"]:
            st.markdown("### 🧭 Navigation")
            
            if st.session_state.current_step != "skill_suggestions":
                if st.button("🎯 Edit Skills", use_container_width=True):
                    st.session_state.current_step = "skill_suggestions"
                    st.rerun()
            
            if st.session_state.current_step != "results":
                if st.button("📄 View Results", use_container_width=True):
                    st.session_state.current_step = "results"
                    st.rerun()

    # Footer
    st.markdown("---")
    st.markdown("*Built with ❤️ by  AI Assistant*")

if __name__ == "__main__":
    main()
